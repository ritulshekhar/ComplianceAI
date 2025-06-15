import io
import re
from typing import Optional
import PyPDF2
import pdfplumber
from docx import Document

class DocumentProcessor:
    """Handle processing of various document formats."""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def process_file(self, uploaded_file) -> str:
        """
        Process uploaded file and extract text content.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
        """
        file_extension = self._get_file_extension(uploaded_file.name)
        
        if file_extension == 'pdf':
            return self._process_pdf(uploaded_file)
        elif file_extension == 'docx':
            return self._process_docx(uploaded_file)
        elif file_extension == 'txt':
            return self._process_txt(uploaded_file)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename."""
        return filename.lower().split('.')[-1] if '.' in filename else ''
    
    def _process_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file."""
        text_content = ""
        
        try:
            # First try with pdfplumber (better for complex layouts)
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            # If pdfplumber didn't extract much text, try PyPDF2
            if len(text_content.strip()) < 100:
                uploaded_file.seek(0)  # Reset file pointer
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
        
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
        
        return self._clean_text(text_content)
    
    def _process_docx(self, uploaded_file) -> str:
        """Extract text from Word document."""
        try:
            doc = Document(uploaded_file)
            
            text_content = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return self._clean_text(text_content)
            
        except Exception as e:
            raise Exception(f"Error processing Word document: {str(e)}")
    
    def _process_txt(self, uploaded_file) -> str:
        """Process plain text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read().decode(encoding)
                    return self._clean_text(content)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            return self._clean_text(content)
            
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newlines
        text = re.sub(r'[ \t]+', ' ', text)       # Multiple spaces/tabs to single space
        
        # Remove special characters that might interfere with analysis
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        return text.strip()
    
    def get_document_stats(self, text: str) -> dict:
        """Get basic statistics about the document."""
        words = len(text.split())
        characters = len(text)
        lines = len(text.split('\n'))
        paragraphs = len([p for p in text.split('\n\n') if p.strip()])
        
        return {
            'word_count': words,
            'character_count': characters,
            'line_count': lines,
            'paragraph_count': paragraphs
        }
    
    def extract_sections(self, text: str) -> dict:
        """Extract common document sections based on headers."""
        sections = {}
        
        # Common section patterns
        section_patterns = [
            r'(?i)^(privacy policy|privacy statement).*?(?=^[A-Z][A-Z\s]+$|\Z)',
            r'(?i)^(terms of service|terms and conditions).*?(?=^[A-Z][A-Z\s]+$|\Z)',
            r'(?i)^(data protection|data handling).*?(?=^[A-Z][A-Z\s]+$|\Z)',
            r'(?i)^(security policy|cybersecurity).*?(?=^[A-Z][A-Z\s]+$|\Z)',
            r'(?i)^(compliance|regulatory).*?(?=^[A-Z][A-Z\s]+$|\Z)'
        ]
        
        for pattern in section_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.DOTALL)
            for match in matches:
                section_name = match.group(1).lower().replace(' ', '_')
                sections[section_name] = match.group(0).strip()
        
        return sections
